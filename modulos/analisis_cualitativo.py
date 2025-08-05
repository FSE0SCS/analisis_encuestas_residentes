# /modulos/analisis_cualitativo.py

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document
from docx.shared import Inches
from collections import Counter
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Descargar recursos de NLTK si no están disponibles
try:
    stopwords.words('spanish')
except nltk.downloader.DownloadError:
    nltk.download('punkt')
    nltk.download('stopwords')

def obtener_preguntas_cualitativas(df):
    """
    Identifica y devuelve las columnas con preguntas de tipo cualitativo.
    """
    columnas = df.columns.tolist()
    # Las preguntas cualitativas son todas las que no están en las primeras 5 columnas
    # ni son las preguntas de tipo cuantitativo
    preguntas_cualitativas_lista = [
        "Comentarios sobre la acogida e integración en el Servicio",
        "¿Realiza tutorías estructuradas cada tres meses con su tutor/a?",
        "¿Conoce los criterios que se aplican para evaluarle de forma continuada? - Informe de evaluación de las rotaciones",
        "¿Qué criterios cree usted que se aplican para evaluarle de forma continuada? - Informe de evaluación de la rotaciones. Selección múltiple",
        "¿Conoce los criterios que se aplican para realizar las evaluaciones anuales? - Informe de evaluación anual del tutor",
        "¿Qué criterios cree usted que se aplican para evaluarle de forma anual? - Informe de evaluación anual del tutor. Selección múltiple",
        "Comentarios sobre la tutorización",
        "¿Dispone de un Libro de Residente?",
        "¿Cómo valora la formación en ética y profesionalismo? - (Comentarios)",
        "¿Cómo valora la planificación y desarrollo de la formación? - (Comentarios)",
        "Comentarios sobre las sesiones clínicas, actividades de investigación y actividades formativas complementarias",
        "Comentarios sobre competencias adquiridas",
        "Comentarios sobre el cumplimiento del calendario y las rotaciones",
        "Comentarios sobre las rotaciones internas",
        "Comentarios sobre las guardias",
        "¿La Unidad Docente le ha informado del Plan de Ayudas del SCS para las Rotaciones Externas?",
        "Comentarios sobre las rotaciones externas",
        "¿Conoce la existencia de la Comisión de Docencia de su Unidad Docente?",
        "¿Conoce al vocal que representa a los residentes en la Comisión de Docencia?",
        "¿Ha planteado alguna vez una queja, propuesta o sugerencia a la Comisión de Docencia?",
        "Comentarios sobre la Comisión de Docencia",
        "¿Le comunican los resultados de la encuesta anual de satisfacción de residentes de su hospital/CCAA?",
        "¿A través de que vía se le comunican dichos resultados? Selección múltiple",
        "Comentarios sobre la comunicación de resultados",
        "Si tuviera que volver a elegir centro para realizar su residencia, ¿volvería a seleccionar este centro?",
        "Comentarios sobre la valoración general",
        "Comentarios de propuestas de mejora"
    ]
    
    preguntas_en_df = [col for col in columnas if col in preguntas_cualitativas_lista]
    
    return preguntas_en_df

def generar_analisis_cualitativo(df, preguntas_seleccionadas):
    """
    Realiza un análisis cualitativo y genera resultados para cada pregunta seleccionada.
    """
    resultados_analisis = []

    for pregunta in preguntas_seleccionadas:
        st.subheader(f"Análisis para la pregunta: {pregunta}")
        
        # Identificar si es una pregunta de comentario o categórica
        if "comentarios" in pregunta.lower():
            st.markdown("### Análisis de Contenido (Pregunta de Texto Libre)")
            
            # Unir todo el texto de los comentarios en una sola cadena
            texto_completo = " ".join(df[pregunta].dropna().astype(str).tolist())
            
            # Limpiar y tokenizar el texto
            palabras = word_tokenize(texto_completo.lower(), language='spanish')
            stop_words = set(stopwords.words('spanish'))
            palabras_limpias = [palabra for palabra in palabras if palabra.isalnum() and palabra not in stop_words]
            
            # Codificación temática y conteo de frecuencias
            frecuencia_palabras = Counter(palabras_limpias)
            palabras_mas_comunes = frecuencia_palabras.most_common(10)
            
            st.write("Frecuencia de palabras clave:")
            df_frecuencia = pd.DataFrame(palabras_mas_comunes, columns=['Palabra Clave', 'Frecuencia'])
            st.dataframe(df_frecuencia)

            # Visualización
            fig, ax = plt.subplots(figsize=(10, 6))
            sns.barplot(x='Frecuencia', y='Palabra Clave', data=df_frecuencia, ax=ax, palette='viridis')
            ax.set_title(f'Top 10 palabras clave para "{pregunta}"')
            st.pyplot(fig)
            plt.close(fig)

            # Generar el prompt para la IA
            prompt = f"Analiza los siguientes comentarios de una encuesta de satisfacción:\n\n{texto_completo}\n\nIdentifica las 5 categorías temáticas más relevantes, el sentimiento general de cada categoría (positivo, negativo, neutro) y extrae las 3 frases más representativas para cada una. Después, resume brevemente los hallazgos principales de este análisis."
            
            resultados_analisis.append({
                "pregunta": pregunta,
                "tipo_analisis": "comentario",
                "df_frecuencia": df_frecuencia,
                "figura": fig,
                "prompt": prompt
            })

        else:
            st.markdown("### Análisis de Frecuencias (Pregunta Categórica)")
            
            # Contar frecuencias
            frecuencias = df[pregunta].value_counts().reset_index()
            frecuencias.columns = ['Respuesta', 'Frecuencia Absoluta']
            frecuencias['Frecuencia Relativa (%)'] = (frecuencias['Frecuencia Absoluta'] / frecuencias['Frecuencia Absoluta'].sum() * 100).round(2)
            
            st.dataframe(frecuencias)
            
            # Gráficos
            fig1, ax1 = plt.subplots(figsize=(10, 6))
            sns.barplot(x='Respuesta', y='Frecuencia Absoluta', data=frecuencias, ax=ax1, palette='viridis')
            ax1.set_title(f'Frecuencia de respuestas para "{pregunta}"')
            st.pyplot(fig1)
            plt.close(fig1)
            
            fig2, ax2 = plt.subplots(figsize=(8, 8))
            ax2.pie(frecuencias['Frecuencia Absoluta'], labels=frecuencias['Respuesta'], autopct='%1.1f%%', startangle=90, colors=sns.color_palette('pastel'))
            ax2.axis('equal')
            ax2.set_title(f'Distribución de respuestas para "{pregunta}"')
            st.pyplot(fig2)
            plt.close(fig2)

            resultados_analisis.append({
                "pregunta": pregunta,
                "tipo_analisis": "categorico",
                "frecuencias": frecuencias,
                "figura_barra": fig1,
                "figura_pastel": fig2,
                "prompt": None
            })
            
    return resultados_analisis

def exportar_analisis_cualitativo_a_word(resultados_analisis):
    """
    Genera un archivo de Word con los resultados del análisis cualitativo.
    """
    document = Document()
    document.add_heading('Resultados del Análisis Cualitativo', 0)
    
    for resultado in resultados_analisis:
        document.add_heading(f"Análisis para: {resultado['pregunta']}", level=1)
        
        if resultado['tipo_analisis'] == 'comentario':
            document.add_paragraph("Análisis de Contenido:")
            document.add_paragraph("Frecuencia de palabras clave:")
            tabla = document.add_table(rows=1, cols=2)
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Palabra Clave'
            hdr_cells[1].text = 'Frecuencia'
            
            for index, row in resultado['df_frecuencia'].iterrows():
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(row['Palabra Clave'])
                row_cells[1].text = str(row['Frecuencia'])
                
            img_buffer = BytesIO()
            resultado['figura'].savefig(img_buffer, format='png')
            img_buffer.seek(0)
            document.add_picture(img_buffer, width=Inches(6))

            document.add_page_break()

        else: # Tipo categórico
            document.add_paragraph("Análisis de Frecuencias:")
            tabla = document.add_table(rows=1, cols=3)
            hdr_cells = tabla.rows[0].cells
            hdr_cells[0].text = 'Respuesta'
            hdr_cells[1].text = 'Frecuencia Absoluta'
            hdr_cells[2].text = 'Frecuencia Relativa (%)'
            
            for index, row in resultado['frecuencias'].iterrows():
                row_cells = tabla.add_row().cells
                row_cells[0].text = str(row['Respuesta'])
                row_cells[1].text = str(row['Frecuencia Absoluta'])
                row_cells[2].text = str(row['Frecuencia Relativa (%)'])
                
            img_buffer_bar = BytesIO()
            resultado['figura_barra'].savefig(img_buffer_bar, format='png')
            img_buffer_bar.seek(0)
            document.add_picture(img_buffer_bar, width=Inches(6))

            img_buffer_pie = BytesIO()
            resultado['figura_pastel'].savefig(img_buffer_pie, format='png')
            img_buffer_pie.seek(0)
            document.add_picture(img_buffer_pie, width=Inches(6))
            
            document.add_page_break()
            
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer